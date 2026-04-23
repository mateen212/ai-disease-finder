"""
Documentation Generator Script
Converts markdown documentation to professionally formatted DOCX file
"""

import re
from pathlib import Path
from datetime import datetime
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx library not installed.")
    print("Install it with: pip install python-docx")
    exit(1)


class DocumentationGenerator:
    """Generate professionally formatted DOCX from markdown"""
    
    def __init__(self, markdown_file, output_file):
        self.markdown_file = markdown_file
        self.output_file = output_file
        self.doc = Document()
        self._setup_styles()
        
    def _setup_styles(self):
        """Configure document styles"""
        # Page setup
        sections = self.doc.sections
        for section in sections:
            section.page_height = Cm(29.7)  # A4
            section.page_width = Cm(21)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
        
        # Title style
        styles = self.doc.styles
        
        # Custom heading styles
        if 'CustomTitle' not in styles:
            title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = 'Arial'
            title_style.font.size = Pt(24)
            title_style.font.bold = True
            title_style.font.color.rgb = RGBColor(0, 0, 128)  # Dark blue
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(12)
        
        # Modify default styles
        for i in range(1, 7):
            style_name = f'Heading {i}'
            if style_name in styles:
                heading = styles[style_name]
                heading.font.name = 'Arial'
                heading.font.color.rgb = RGBColor(0, 0, 102)  # Dark blue
                heading.paragraph_format.space_before = Pt(12 if i <= 2 else 6)
                heading.paragraph_format.space_after = Pt(6)
                heading.paragraph_format.keep_with_next = True
        
        # Body text
        if 'Normal' in styles:
            normal = styles['Normal']
            normal.font.name = 'Times New Roman'
            normal.font.size = Pt(11)
            normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            normal.paragraph_format.space_after = Pt(6)
            normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    def _add_title_page(self):
        """Create professional title page"""
        # Title
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("Hybrid Neuro-Symbolic Clinical Decision Support System")
        run.font.name = 'Arial'
        run.font.size = Pt(26)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 102)
        
        # Subtitle
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("Multi-Disease Diagnosis with Explainable AI")
        run.font.name = 'Arial'
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0, 0, 128)
        
        # Spacing
        for _ in range(3):
            self.doc.add_paragraph()
        
        # Author info
        self._add_centered_text("System Documentation", 14, True)
        self.doc.add_paragraph()
        self._add_centered_text(f"Version 1.0", 12, False)
        self._add_centered_text(f"Date: {datetime.now().strftime('%B %d, %Y')}", 12, False)
        
        # Spacing
        for _ in range(5):
            self.doc.add_paragraph()
        
        # Institution
        self._add_centered_text("Medical AI Research Team", 14, True)
        self._add_centered_text("Knowledge-Based Systems Project", 12, False)
        
        # Page break
        self.doc.add_page_break()
    
    def _add_centered_text(self, text, size, bold):
        """Add centered text with formatting"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(size)
        run.font.bold = bold
    
    def _parse_markdown_line(self, line):
        """Determine line type and content"""
        line = line.rstrip()
        
        # Headers
        if line.startswith('# '):
            return 'h1', line[2:]
        elif line.startswith('## '):
            return 'h2', line[3:]
        elif line.startswith('### '):
            return 'h3', line[4:]
        elif line.startswith('#### '):
            return 'h4', line[5:]
        elif line.startswith('##### '):
            return 'h5', line[6:]
        elif line.startswith('###### '):
            return 'h6', line[7:]
        
        # Bold text
        elif line.startswith('**') and line.endswith('**'):
            return 'bold', line[2:-2]
        
        # Lists
        elif line.startswith('- ') or line.startswith('* '):
            return 'bullet', line[2:]
        elif re.match(r'^\d+\.\s', line):
            match = re.match(r'^(\d+)\.\s(.+)$', line)
            if match:
                return 'numbered', match.group(2)
        
        # Code blocks
        elif line.startswith('```'):
            return 'code_marker', line[3:]
        
        # Horizontal rule
        elif line in ['---', '***', '___']:
            return 'hr', ''
        
        # Empty line
        elif line.strip() == '':
            return 'empty', ''
        
        # Regular paragraph
        else:
            return 'text', line
    
    def _add_table_from_markdown(self, lines, start_idx):
        """Parse and add markdown table"""
        table_lines = []
        idx = start_idx
        
        # Collect table lines
        while idx < len(lines) and '|' in lines[idx]:
            table_lines.append(lines[idx])
            idx += 1
        
        if len(table_lines) < 2:
            return idx
        
        # Parse headers
        headers = [cell.strip() for cell in table_lines[0].split('|') if cell.strip()]
        
        # Skip separator line
        if len(table_lines) > 1 and '---' in table_lines[1]:
            data_start = 2
        else:
            data_start = 1
        
        # Parse data rows
        rows = []
        for i in range(data_start, len(table_lines)):
            cells = [cell.strip() for cell in table_lines[i].split('|') if cell.strip()]
            if cells:
                rows.append(cells)
        
        # Create table
        if rows:
            table = self.doc.add_table(rows=len(rows) + 1, cols=len(headers))
            table.style = 'Light Grid Accent 1'
            
            # Add headers
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                if i < len(header_cells):
                    header_cells[i].text = header
                    # Bold header
                    for paragraph in header_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(10)
            
            # Add data
            for row_idx, row_data in enumerate(rows):
                cells = table.rows[row_idx + 1].cells
                for col_idx, cell_data in enumerate(row_data):
                    if col_idx < len(cells):
                        cells[col_idx].text = cell_data
                        # Format cell text
                        for paragraph in cells[col_idx].paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(9)
        
        return idx
    
    def generate(self):
        """Generate DOCX document from markdown"""
        print(f"Reading markdown file: {self.markdown_file}")
        
        with open(self.markdown_file, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        print("Generating professional DOCX document...")
        
        # Add title page
        self._add_title_page()
        
        # Process content
        in_code_block = False
        code_content = []
        i = 0
        
        while i < len(lines):
            line = lines[i]
            line_type, content = self._parse_markdown_line(line)
            
            # Handle code blocks
            if line_type == 'code_marker':
                if in_code_block:
                    # End of code block
                    if code_content:
                        code_para = self.doc.add_paragraph('\n'.join(code_content))
                        code_para.style = 'Normal'
                        for run in code_para.runs:
                            run.font.name = 'Courier New'
                            run.font.size = Pt(9)
                        # Add light gray background
                        shading_elm = OxmlElement('w:shd')
                        shading_elm.set(qn('w:fill'), 'F0F0F0')
                        code_para._element.get_or_add_pPr().append(shading_elm)
                    code_content = []
                    in_code_block = False
                else:
                    in_code_block = True
                i += 1
                continue
            
            if in_code_block:
                code_content.append(line.rstrip())
                i += 1
                continue
            
            # Handle different line types
            if line_type == 'h1':
                p = self.doc.add_paragraph(content, style='Heading 1')
                p.runs[0].font.size = Pt(20)
            elif line_type == 'h2':
                p = self.doc.add_paragraph(content, style='Heading 2')
                p.runs[0].font.size = Pt(16)
            elif line_type == 'h3':
                p = self.doc.add_paragraph(content, style='Heading 3')
                p.runs[0].font.size = Pt(14)
            elif line_type == 'h4':
                p = self.doc.add_paragraph(content, style='Heading 4')
                p.runs[0].font.size = Pt(12)
            elif line_type == 'h5':
                p = self.doc.add_paragraph(content, style='Heading 5')
            elif line_type == 'h6':
                p = self.doc.add_paragraph(content, style='Heading 6')
            elif line_type == 'bold':
                p = self.doc.add_paragraph()
                run = p.add_run(content)
                run.font.bold = True
            elif line_type == 'bullet':
                p = self.doc.add_paragraph(content, style='List Bullet')
            elif line_type == 'numbered':
                p = self.doc.add_paragraph(content, style='List Number')
            elif line_type == 'hr':
                # Add horizontal line
                p = self.doc.add_paragraph()
                p.paragraph_format.border_bottom = True
            elif line_type == 'empty':
                # Don't add too many empty paragraphs
                if i > 0 and i < len(lines) - 1:
                    self.doc.add_paragraph()
            elif line_type == 'text':
                # Check if next lines form a table
                if '|' in line:
                    i = self._add_table_from_markdown(lines, i)
                    continue
                else:
                    # Process inline formatting
                    content = self._process_inline_formatting(content)
                    if content.strip():
                        self.doc.add_paragraph(content)
            
            i += 1
        
        # Save document
        print(f"Saving document to: {self.output_file}")
        self.doc.save(self.output_file)
        print(f"✓ Documentation generated successfully!")
        print(f"  File: {self.output_file}")
        print(f"  Size: {Path(self.output_file).stat().st_size / 1024:.1f} KB")
    
    def _process_inline_formatting(self, text):
        """Process inline markdown formatting (bold, italic, code)"""
        # Bold: **text** or __text__
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'__(.+?)__', r'\1', text)
        
        # Italic: *text* or _text_
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        text = re.sub(r'_(.+?)_', r'\1', text)
        
        # Inline code: `code`
        text = re.sub(r'`(.+?)`', r'\1', text)
        
        # Links: [text](url)
        text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
        
        return text


def main():
    """Main execution"""
    print("="*60)
    print("Professional Documentation Generator")
    print("="*60)
    print()
    
    # File paths
    markdown_file = Path(__file__).parent / "COMPREHENSIVE_SYSTEM_DOCUMENTATION.md"
    output_file = Path(__file__).parent / "System_Documentation_Complete.docx"
    
    if not markdown_file.exists():
        print(f"ERROR: Markdown file not found: {markdown_file}")
        return
    
    # Generate documentation
    generator = DocumentationGenerator(markdown_file, output_file)
    generator.generate()
    
    print()
    print("="*60)
    print("Documentation generation complete!")
    print("="*60)
    print()
    print("The professionally formatted Word document has been created.")
    print(f"Location: {output_file.absolute()}")
    print()
    print("You can now:")
    print("  1. Open it in Microsoft Word or LibreOffice")
    print("  2. Make final adjustments to formatting if needed")
    print("  3. Export to PDF if required")
    print()


if __name__ == "__main__":
    main()
